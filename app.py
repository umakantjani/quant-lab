"""
QUANT LAB - MODERN FINANCIAL TERMINAL
A sophisticated quantitative analysis platform with React-inspired UI architecture
"""

import streamlit as st
import pandas as pd
import os
import subprocess
import sys
import time
from pathlib import Path
from sqlalchemy import text
from logic.db_config import get_engine
from fpdf import FPDF
from docx import Document
from io import BytesIO

# ============================================================================
# CONFIGURATION
# ============================================================================

st.set_page_config(
    page_title="QuantLab Terminal",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://github.com/umakantjani/quant-lab',
        'Report a bug': 'https://github.com/umakantjani/quant-lab/issues',
        'About': "QuantLab Terminal - Professional Quantitative Analysis Platform"
    }
)

LOGIC_DIR = "logic"

# ============================================================================
# STYLE INJECTION
# ============================================================================

def load_css():
    """Load external CSS file"""
    css_file = Path(__file__).parent / "styles.css"
    
    if css_file.exists():
        with open(css_file) as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    else:
        # Fallback to inline CSS if file not found
        st.markdown("""
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Archivo:wght@300;400;600;700&family=IBM+Plex+Mono:wght@400;600&family=Sora:wght@600;700;800&display=swap');
            
            :root {
                --color-bg-primary: #0a0e1a;
                --color-accent-primary: #00ff88;
                --font-display: 'Sora', sans-serif;
                --font-body: 'Archivo', sans-serif;
                --font-mono: 'IBM Plex Mono', monospace;
            }
            
            .stApp { background: var(--color-bg-primary); }
            h1, h2, h3 { font-family: var(--font-display); }
            p, div { font-family: var(--font-body); }
        </style>
        """, unsafe_allow_html=True)

load_css()

# ============================================================================
# CUSTOM COMPONENTS (React-style)
# ============================================================================

class Components:
    """Reusable UI components"""
    
    @staticmethod
    def terminal_header(title: str, subtitle: str = ""):
        """Terminal-style page header"""
        st.markdown(f"""
        <div style="margin-bottom: 2rem;">
            <h1 style="margin-bottom: 0.5rem;">{title}</h1>
            {f'<p style="font-size: 1.1rem; color: #a8b2c0;">{subtitle}</p>' if subtitle else ''}
        </div>
        """, unsafe_allow_html=True)
    
    @staticmethod
    def metric_card(label: str, value: str, delta: str = None, icon: str = "📊"):
        """Modern metric card"""
        delta_html = f'<div style="font-size: 0.9rem; color: #00ff88;">{delta}</div>' if delta else ''
        
        st.markdown(f"""
        <div style="
            background: linear-gradient(135deg, rgba(0, 255, 136, 0.05), rgba(0, 136, 255, 0.05));
            padding: 1.5rem;
            border-radius: 12px;
            border: 1px solid rgba(255, 255, 255, 0.08);
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.4);
        ">
            <div style="font-size: 0.85rem; color: #6b7480; text-transform: uppercase; letter-spacing: 0.05em; margin-bottom: 0.5rem;">
                {icon} {label}
            </div>
            <div style="font-size: 2rem; font-weight: 700; color: #00ff88; font-family: 'IBM Plex Mono', monospace;">
                {value}
            </div>
            {delta_html}
        </div>
        """, unsafe_allow_html=True)
    
    @staticmethod
    def action_button_group(buttons: list):
        """Group of action buttons"""
        cols = st.columns(len(buttons))
        for col, (label, icon, key) in zip(cols, buttons):
            with col:
                if st.button(f"{icon} {label}", key=key, use_container_width=True):
                    return key
        return None
    
    @staticmethod
    def status_badge(status: str, text: str = None):
        """Colored status badge"""
        colors = {
            "success": "#00ff88",
            "warning": "#ffaa00",
            "danger": "#ff3366",
            "info": "#0088ff",
            "neutral": "#6b7480"
        }
        color = colors.get(status, colors["neutral"])
        display_text = text or status.upper()
        
        return f'<span style="background: {color}20; color: {color}; padding: 0.25rem 0.75rem; border-radius: 4px; font-size: 0.85rem; font-weight: 600;">{display_text}</span>'
    
    @staticmethod
    def section_divider(text: str = ""):
        """Visual section separator"""
        if text:
            st.markdown(f"""
            <div style="display: flex; align-items: center; margin: 2rem 0;">
                <div style="flex: 1; height: 1px; background: linear-gradient(to right, transparent, rgba(255,255,255,0.1), transparent);"></div>
                <div style="margin: 0 1rem; color: #6b7480; font-size: 0.9rem; text-transform: uppercase; letter-spacing: 0.1em;">{text}</div>
                <div style="flex: 1; height: 1px; background: linear-gradient(to right, transparent, rgba(255,255,255,0.1), transparent);"></div>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown('<div style="height: 1px; background: rgba(255,255,255,0.05); margin: 2rem 0;"></div>', unsafe_allow_html=True)

# ============================================================================
# CORE FUNCTIONALITY
# ============================================================================

@st.cache_resource
def get_db_connection():
    """Cached database connection"""
    try:
        return get_engine()
    except Exception as e:
        st.error(f"Database connection failed: {e}")
        return None

engine = get_db_connection()

def load_data(query: str) -> pd.DataFrame:
    """Safe data loading with error handling"""
    if engine is None:
        return pd.DataFrame()
    try:
        return pd.read_sql(query, engine)
    except Exception as e:
        st.error(f"Data load error: {e}")
        return pd.DataFrame()

def run_logic_script(script_name: str) -> bool:
    """Execute backend script with live progress"""
    script_path = os.path.join(LOGIC_DIR, script_name)
    
    if not os.path.exists(script_path):
        st.error(f"❌ Script not found: {script_path}")
        return False
    
    with st.status(f"🚀 Running {script_name}...", expanded=True) as status:
        env = os.environ.copy()
        env["PYTHONPATH"] = os.getcwd()
        
        logs = []
        process = subprocess.Popen(
            [sys.executable, script_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            env=env,
            bufsize=1
        )
        
        log_container = st.empty()
        
        while True:
            line = process.stdout.readline()
            if not line and process.poll() is not None:
                break
            if line:
                logs.append(line)
                log_container.code("".join(logs[-20:]), language="bash")
        
        if process.returncode == 0:
            status.update(label=f"✅ {script_name} Complete", state="complete")
            return True
        else:
            status.update(label=f"❌ {script_name} Failed", state="error")
            return False

# ============================================================================
# EXPORT FUNCTIONS
# ============================================================================

def create_pdf(text_content: str, ticker: str) -> BytesIO:
    """Generate PDF report"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    clean_text = text_content.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 10, clean_text)
    
    buffer = BytesIO()
    pdf_output = pdf.output(dest='S').encode('latin-1')
    buffer.write(pdf_output)
    buffer.seek(0)
    return buffer

def create_docx(text_content: str, ticker: str) -> BytesIO:
    """Generate Word document"""
    doc = Document()
    doc.add_heading(f'Investment Report: {ticker}', 0)
    for line in text_content.split('\n'):
        doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================================================
# DOCUMENTATION
# ============================================================================

def render_guide(topic: str):
    """Display contextual help"""
    guides = {
        "valuation": {
            "title": "📘 Valuation Guide",
            "content": """
            **The Philosophy:** Calculate the *Intrinsic Value* using a **Discounted Cash Flow (DCF)** model.
            
            - **Upside %:** Gap between Price and Value. Look for > 15%.
            - **WACC:** The "Hurdle Rate." Higher = Riskier business.
            - **Rating:** Synthetic credit rating based on financial health.
            """
        },
        "technicals": {
            "title": "📈 Technical Analysis Guide",
            "content": """
            **The Sniper's Trinity:** Wait for three aligned conditions.
            
            - **1. The Tide (Trend):** Price > 200 SMA (Bullish market)
            - **2. The Zone (Setup):** RSI < 30 (Oversold condition)
            - **3. The Trigger (Action):** Hammer or Engulfing candle pattern
            """
        }
    }
    
    if topic in guides:
        with st.expander(guides[topic]["title"]):
            st.markdown(guides[topic]["content"])

# ============================================================================
# SIDEBAR NAVIGATION
# ============================================================================

def render_sidebar():
    """Modern sidebar with navigation"""
    with st.sidebar:
        st.markdown("""
        <div style="text-align: center; padding: 2rem 0; border-bottom: 2px solid rgba(255,255,255,0.08);">
            <div style="font-family: 'IBM Plex Mono', monospace; font-size: 1.5rem; font-weight: 800; letter-spacing: 0.1em; color: #00ff88;">
                QUANT LAB
            </div>
            <div style="font-size: 0.75rem; color: #6b7480; margin-top: 0.5rem; letter-spacing: 0.15em;">
                TERMINAL v2.0
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        mode = st.radio(
            "NAVIGATION",
            [
                "⚡ Daily Alpha",
                "🔬 Research Lab",
                "🧬 AI Reports",
                "⚙️ System Admin"
            ],
            label_visibility="collapsed"
        )
        
        Components.section_divider()
        
        # Status indicators
        st.markdown("**SYSTEM STATUS**")
        if engine:
            st.markdown(Components.status_badge("success", "Database Connected"), unsafe_allow_html=True)
        else:
            st.markdown(Components.status_badge("danger", "Database Error"), unsafe_allow_html=True)
        
        return mode

# ============================================================================
# PAGE: DAILY ALPHA
# ============================================================================

def page_daily_alpha():
    """Daily action generator page"""
    Components.terminal_header(
        "⚡ Daily Alpha Generator",
        "Generate today's buy list based on Deep Value and Price Momentum"
    )
    
    # Action buttons
    st.markdown("### Pipeline Control")
    action = Components.action_button_group([
        ("Scanner", "🔍", "scan"),
        ("Valuation", "💎", "val"),
        ("Technicals", "📈", "tech"),
        ("Hedges", "🛡", "hedge"),
        ("Orders", "🛒", "orders")
    ])
    
    if action:
        script_map = {
            "scan": "scanner.py",
            "val": "valuation.py",
            "tech": "technical_models.py",
            "hedge": "etf_mapper.py",
            "orders": "orders.py"
        }
        run_logic_script(script_map[action])
    
    Components.section_divider("Results Dashboard")
    
    # Tabs for results
    tab_val, tab_tech, tab_etf, tab_ord = st.tabs([
        "💎 Valuation Report",
        "📈 Technical Signals",
        "🛡 ETF Strategy",
        "🛒 Buy Orders"
    ])
    
    with tab_val:
        render_guide("valuation")
        df = load_data("SELECT * FROM alpha_valuation ORDER BY upside_pct DESC")
        
        if not df.empty:
            col1, col2, col3 = st.columns(3)
            with col1:
                Components.metric_card(
                    "Opportunities",
                    str(len(df)),
                    icon="💎"
                )
            with col2:
                avg_upside = df['upside_pct'].mean()
                Components.metric_card(
                    "Avg Upside",
                    f"{avg_upside:.1f}%",
                    icon="📊"
                )
            with col3:
                best_pick = df.iloc[0]['ticker']
                Components.metric_card(
                    "Top Pick",
                    best_pick,
                    icon="🎯"
                )
            
            st.markdown("<br>", unsafe_allow_html=True)
            st.dataframe(
                df.style.format({
                    "current_price": "${:.2f}",
                    "intrinsic_value": "${:.2f}",
                    "upside_pct": "{:.1f}%",
                    "wacc_pct": "{:.1f}%"
                }).background_gradient(subset=['upside_pct'], cmap="Greens"),
                use_container_width=True,
                height=500
            )
        else:
            st.info("📊 No valuation data available. Run the pipeline first.")
    
    with tab_tech:
        render_guide("technicals")
        df_tech = load_data('SELECT * FROM technical_signals WHERE "Signal" != \'WAIT\' ORDER BY "Score" DESC')
        
        if not df_tech.empty:
            st.dataframe(
                df_tech.style.map(
                    lambda x: 'color: #00ff88' if 'BUY' in str(x) else ('color: #ff3366' if 'SELL' in str(x) else ''),
                    subset=['Signal']
                ),
                use_container_width=True
            )
        else:
            st.info("📈 No active technical signals.")
    
    with tab_etf:
        df_etf = load_data("SELECT * FROM etf_hedges ORDER BY weight DESC")
        
        if not df_etf.empty:
            st.dataframe(
                df_etf.style.format({"weight": "{:.1%}"}).background_gradient(subset=['weight'], cmap="Reds"),
                use_container_width=True
            )
        else:
            st.info("🛡 No ETF hedge data available.")
    
    with tab_ord:
        df_ord = load_data("SELECT * FROM alpha_orders")
        
        if not df_ord.empty:
            st.dataframe(
                df_ord.style.format({"Limit_Price": "${:.2f}", "Est_Value": "${:,.2f}"}),
                use_container_width=True
            )
            
            csv = df_ord.to_csv(index=False).encode('utf-8')
            st.download_button(
                "⬇️ Download Order File (CSV)",
                csv,
                "basket_orders.csv",
                "text/csv",
                use_container_width=True
            )
        else:
            st.info("🛒 No orders generated yet.")

# ============================================================================
# PAGE: RESEARCH LAB
# ============================================================================

def page_research_lab():
    """Deep research and analysis page"""
    Components.terminal_header(
        "🔬 Quantitative Research Lab",
        "Advanced factor analysis and technical screening"
    )
    
    c1, c2 = st.columns(2)
    with c1:
        if st.button("🏆 Run Factor Scoring", use_container_width=True):
            run_logic_script("model_engine.py")
    with c2:
        if st.button("🎯 Run Sniper Scan", use_container_width=True):
            run_logic_script("technical_models.py")
    
    Components.section_divider("Analysis Results")
    
    t1, t2 = st.tabs(["🏆 Factor Rankings", "🎯 Sniper Signals"])
    
    with t1:
        df_rank = load_data('SELECT * FROM quant_rankings ORDER BY "TOTAL_SCORE" DESC')
        
        if not df_rank.empty:
            st.dataframe(
                df_rank.style.background_gradient(subset=['TOTAL_SCORE'], cmap="RdYlGn"),
                use_container_width=True,
                height=600
            )
        else:
            st.info("🏆 No ranking data available.")
    
    with t2:
        render_guide("technicals")
        
        filter_mode = st.radio("Display Mode:", ["Active Signals Only", "Full Checklist"], horizontal=True)
        
        df_tech = load_data('SELECT * FROM technical_signals ORDER BY "Score" DESC')
        
        if not df_tech.empty:
            if filter_mode == "Active Signals Only":
                df_display_source = df_tech[df_tech['Signal'] != 'WAIT']
            else:
                df_display_source = df_tech
            
            df_display = pd.DataFrame({
                'Ticker': df_display_source['ticker'],
                'Price': df_display_source['price'].apply(lambda x: f"${x:.2f}"),
                'Score': df_display_source['Score'],
                'Signal': df_display_source['Signal'],
                '🌊 Trend': df_display_source.apply(
                    lambda x: f"✅ BULL (${x['trend_200_sma']:.2f})" if x['trend_pass'] else f"❌ BEAR (${x['trend_200_sma']:.2f})",
                    axis=1
                ),
                '📉 Zone': df_display_source.apply(
                    lambda x: f"✅ LOW ({x['rsi']})" if x['rsi'] < 35 else f"⚠️ ({x['rsi']})",
                    axis=1
                ),
                '🕯️ Trigger': df_display_source['trigger_type'].apply(
                    lambda x: f"✅ {x}" if x != "None" else "❌ Wait"
                )
            })
            
            st.dataframe(
                df_display.style.background_gradient(subset=['Score'], cmap="RdYlGn").map(
                    lambda x: 'color: #00ff88' if 'STRONG' in str(x) else '',
                    subset=['Signal']
                ),
                use_container_width=True
            )
        else:
            st.info("🎯 No technical data available.")

# ============================================================================
# PAGE: AI REPORTS
# ============================================================================

def page_ai_reports():
    """AI-powered research reports"""
    Components.terminal_header(
        "🧬 Asset Lab Report",
        "Automated institutional research powered by Gemini AI"
    )
    
    try:
        from logic.report_engine import generate_ai_report, save_report, get_report_history
    except ImportError:
        st.error("❌ Report engine not available. Check logic/report_engine.py")
        return
    
    # Session state management
    if "current_report" not in st.session_state:
        st.session_state.current_report = ""
    if "current_ticker" not in st.session_state:
        st.session_state.current_ticker = ""
    
    tab_gen, tab_hist = st.tabs(["✨ Generate New Report", "📜 Report Archive"])
    
    with tab_gen:
        c1, c2 = st.columns([1, 3])
        
        with c1:
            st.markdown("### New Analysis")
            
            db_tickers = load_data("SELECT DISTINCT ticker FROM prices ORDER BY ticker")['ticker'].tolist()
            if not db_tickers:
                db_tickers = ["AAPL", "MSFT", "TSLA", "GOOGL", "NVDA"]
            
            ticker_input = st.selectbox("Select Ticker", db_tickers)
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            if st.button("✨ Generate Report", use_container_width=True, type="primary"):
                with st.status(f"🚀 Analyzing {ticker_input}...", expanded=True) as status:
                    st.write("📡 Fetching market data...")
                    time.sleep(0.5)
                    st.write("🧠 Synthesizing narrative...")
                    
                    report = generate_ai_report(ticker_input)
                    st.session_state.current_report = report
                    st.session_state.current_ticker = ticker_input
                    
                    status.update(label="✅ Analysis Complete", state="complete", expanded=False)
        
        with c2:
            if st.session_state.current_report:
                st.markdown(st.session_state.current_report)
                
                Components.section_divider()
                
                # Save and export controls
                sc1, sc2, sc3 = st.columns(3)
                
                with sc1:
                    if st.button("💾 Save to Archive", use_container_width=True):
                        version = save_report(st.session_state.current_ticker, st.session_state.current_report)
                        st.success(f"✅ Saved to database (v{version})")
                
                with sc2:
                    docx_file = create_docx(st.session_state.current_report, st.session_state.current_ticker)
                    st.download_button(
                        "📄 Download Word",
                        docx_file,
                        f"{st.session_state.current_ticker}_Report.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                with sc3:
                    pdf_file = create_pdf(st.session_state.current_report, st.session_state.current_ticker)
                    st.download_button(
                        "📕 Download PDF",
                        pdf_file,
                        f"{st.session_state.current_ticker}_Report.pdf",
                        "application/pdf",
                        use_container_width=True
                    )
            else:
                st.info("👈 Select a ticker and generate a report to get started")
    
    with tab_hist:
        st.markdown("### 📜 Research History")
        
        df_hist = get_report_history()
        
        if df_hist.empty:
            st.info("No saved reports yet. Generate and save one!")
        else:
            hist_tickers = df_hist['ticker'].unique().tolist()
            sel_hist_ticker = st.selectbox("Filter by Ticker", ["All"] + hist_tickers)
            
            filtered_hist = get_report_history(None if sel_hist_ticker == "All" else sel_hist_ticker)
            
            st.dataframe(
                filtered_hist[['report_date', 'ticker', 'version', 'signal', 'created_at']],
                use_container_width=True,
                hide_index=True
            )
            
            if not filtered_hist.empty:
                Components.section_divider("Report Viewer")
                
                report_id = st.selectbox(
                    "Select Report to View",
                    filtered_hist.index,
                    format_func=lambda x: f"{filtered_hist.loc[x, 'report_date']} - {filtered_hist.loc[x, 'ticker']} (v{filtered_hist.loc[x, 'version']})"
                )
                
                if report_id is not None:
                    st.markdown(f"**{filtered_hist.loc[report_id, 'ticker']} (v{filtered_hist.loc[report_id, 'version']})**")
                    st.markdown(filtered_hist.loc[report_id, 'content'])

# ============================================================================
# PAGE: SYSTEM ADMIN
# ============================================================================

def page_system_admin():
    """System administration and maintenance"""
    Components.terminal_header(
        "⚙️ System Administration",
        "Database maintenance and system diagnostics"
    )
    
    st.warning("⚠️ **Restricted Area:** These operations affect the entire system")
    
    with st.expander("🗄️ DATABASE MAINTENANCE", expanded=True):
        c1, c2, c3 = st.columns(3)
        
        with c1:
            if st.button("🔄 Reset Database", use_container_width=True):
                if st.button("⚠️ Confirm Reset", use_container_width=True):
                    run_logic_script("setup_db.py")
        
        with c2:
            if st.button("📥 Ingest Tickers", use_container_width=True):
                run_logic_script("ingest_tickers.py")
        
        with c3:
            if st.button("🔄 Full Data Update", use_container_width=True):
                run_logic_script("data_pipeline.py")
    
    with st.expander("🔍 SYSTEM DIAGNOSTICS"):
        if st.button("🏥 Run Health Check", use_container_width=True):
            run_logic_script("diagnose.py")
        
        # Display system info
        if engine:
            st.success("✅ Database connection active")
            
            # Table stats
            tables = ['tickers', 'prices', 'fundamentals', 'alpha_valuation', 'technical_signals']
            
            st.markdown("### Table Statistics")
            
            for table in tables:
                try:
                    count_df = load_data(f"SELECT COUNT(*) as count FROM {table}")
                    if not count_df.empty:
                        count = count_df.iloc[0]['count']
                        st.metric(f"{table.upper()}", f"{count:,} records")
                except:
                    pass

# ============================================================================
# MAIN APPLICATION ROUTER
# ============================================================================

def main():
    """Main application entry point"""
    
    # Render sidebar and get navigation selection
    mode = render_sidebar()
    
    # Route to appropriate page
    if mode == "⚡ Daily Alpha":
        page_daily_alpha()
    
    elif mode == "🔬 Research Lab":
        page_research_lab()
    
    elif mode == "🧬 AI Reports":
        page_ai_reports()
    
    elif mode == "⚙️ System Admin":
        page_system_admin()

# ============================================================================
# APPLICATION ENTRY POINT
# ============================================================================

if __name__ == "__main__":
    main()
