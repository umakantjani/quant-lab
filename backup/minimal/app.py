"""
QUANT LAB - MINIMAL DESIGN
Inspired by Jony Ive's design philosophy: "Simplicity is the ultimate sophistication"

Data is king. Everything else whispers.
"""

import streamlit as st
import pandas as pd
import os
import subprocess
import sys
import time
from pathlib import Path
from sqlalchemy import text
from logic.db_config import get_engine
from fpdf import FPDF
from docx import Document
from io import BytesIO

# ============================================================================
# CONFIGURATION
# ============================================================================

st.set_page_config(
    page_title="QuantLab",
    page_icon="□",
    layout="wide",
    initial_sidebar_state="expanded"
)

LOGIC_DIR = "logic"

# ============================================================================
# STYLE INJECTION - MINIMAL APPLE AESTHETIC
# ============================================================================

def load_css():
    """Load minimal Apple-inspired CSS"""
    css_file = Path(__file__).parent / "styles_minimal.css"
    
    if css_file.exists():
        with open(css_file) as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

load_css()

# ============================================================================
# MINIMAL COMPONENTS - DATA FIRST
# ============================================================================

class UI:
    """Minimal UI components - Apple style"""
    
    @staticmethod
    def header(title: str, subtitle: str = ""):
        """Minimal page header"""
        st.markdown(f"## {title}")
        if subtitle:
            st.markdown(f'<p class="text-secondary" style="margin-top: -8px; margin-bottom: 24px;">{subtitle}</p>', 
                       unsafe_allow_html=True)
    
    @staticmethod
    def divider():
        """Subtle divider"""
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    
    @staticmethod
    def badge(status: str, text: str = None):
        """Minimal status badge"""
        badge_map = {
            "success": "badge-success",
            "danger": "badge-danger",
            "warning": "badge-warning",
            "neutral": "badge-neutral"
        }
        css_class = badge_map.get(status, "badge-neutral")
        display_text = text or status.upper()
        
        return f'<span class="badge {css_class}">{display_text}</span>'
    
    @staticmethod
    def small_metric(label: str, value: str):
        """Compact metric display"""
        st.markdown(f"""
        <div style="margin-bottom: 16px;">
            <div class="label-text" style="margin-bottom: 4px;">{label}</div>
            <div style="font-family: var(--font-mono); font-size: 1.25rem; color: var(--color-text-primary);">
                {value}
            </div>
        </div>
        """, unsafe_allow_html=True)

# ============================================================================
# CORE FUNCTIONALITY
# ============================================================================

@st.cache_resource
def get_db_connection():
    """Cached database connection"""
    try:
        return get_engine()
    except Exception as e:
        return None

engine = get_db_connection()

def load_data(query: str) -> pd.DataFrame:
    """Load data from database"""
    if engine is None:
        return pd.DataFrame()
    try:
        return pd.read_sql(query, engine)
    except:
        return pd.DataFrame()

def run_logic_script(script_name: str) -> bool:
    """Execute backend script"""
    script_path = os.path.join(LOGIC_DIR, script_name)
    
    if not os.path.exists(script_path):
        st.error(f"Script not found: {script_path}")
        return False
    
    with st.status(f"Running {script_name.replace('.py', '')}...", expanded=False) as status:
        env = os.environ.copy()
        env["PYTHONPATH"] = os.getcwd()
        
        process = subprocess.Popen(
            [sys.executable, script_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            env=env,
            bufsize=1
        )
        
        while process.poll() is None:
            time.sleep(0.1)
        
        if process.returncode == 0:
            status.update(label=f"✓ {script_name.replace('.py', '')} complete", state="complete")
            return True
        else:
            status.update(label=f"✗ {script_name.replace('.py', '')} failed", state="error")
            return False

# ============================================================================
# EXPORT FUNCTIONS
# ============================================================================

def create_pdf(text_content: str, ticker: str) -> BytesIO:
    """Generate PDF report"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    clean_text = text_content.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 8, clean_text)
    
    buffer = BytesIO()
    pdf_output = pdf.output(dest='S').encode('latin-1')
    buffer.write(pdf_output)
    buffer.seek(0)
    return buffer

def create_docx(text_content: str, ticker: str) -> BytesIO:
    """Generate Word document"""
    doc = Document()
    doc.add_heading(f'{ticker} Report', 0)
    for line in text_content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================================================
# SIDEBAR - MINIMAL NAVIGATION
# ============================================================================

def render_sidebar():
    """Minimal sidebar navigation"""
    with st.sidebar:
        st.markdown("# QuantLab")
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        mode = st.radio(
            "",
            [
                "Daily Alpha",
                "Research",
                "Reports",
                "System"
            ],
            label_visibility="collapsed"
        )
        
        UI.divider()
        
        # Minimal status
        st.markdown('<p class="label-text">Status</p>', unsafe_allow_html=True)
        if engine:
            st.markdown(UI.badge("success", "Connected"), unsafe_allow_html=True)
        else:
            st.markdown(UI.badge("danger", "Error"), unsafe_allow_html=True)
        
        return mode

# ============================================================================
# PAGE: DAILY ALPHA
# ============================================================================

def page_daily_alpha():
    """Daily alpha generation - data focused"""
    UI.header("Daily Alpha", "Quantitative opportunities for today")
    
    # Compact controls
    cols = st.columns(5)
    with cols[0]:
        if st.button("Scanner", use_container_width=True):
            run_logic_script("scanner.py")
    with cols[1]:
        if st.button("Valuation", use_container_width=True):
            run_logic_script("valuation.py")
    with cols[2]:
        if st.button("Technicals", use_container_width=True):
            run_logic_script("technical_models.py")
    with cols[3]:
        if st.button("Hedges", use_container_width=True):
            run_logic_script("etf_mapper.py")
    with cols[4]:
        if st.button("Orders", use_container_width=True):
            run_logic_script("orders.py")
    
    UI.divider()
    
    # Tabs - data first
    tab1, tab2, tab3, tab4 = st.tabs(["Valuation", "Technicals", "Hedges", "Orders"])
    
    with tab1:
        df = load_data("SELECT * FROM alpha_valuation ORDER BY upside_pct DESC")
        
        if not df.empty:
            # Compact summary metrics
            c1, c2, c3 = st.columns(3)
            with c1:
                UI.small_metric("Opportunities", str(len(df)))
            with c2:
                UI.small_metric("Avg Upside", f"{df['upside_pct'].mean():.1f}%")
            with c3:
                UI.small_metric("Top Pick", df.iloc[0]['ticker'])
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            # Data table - no distractions
            st.dataframe(
                df[[
                    'ticker', 'current_price', 'intrinsic_value', 
                    'upside_pct', 'wacc_pct', 'synthetic_spread'
                ]].style.format({
                    "current_price": "${:.2f}",
                    "intrinsic_value": "${:.2f}",
                    "upside_pct": "{:.1f}%",
                    "wacc_pct": "{:.1f}%",
                    "synthetic_spread": "{:.2f}%"
                }),
                use_container_width=True,
                height=500,
                hide_index=True
            )
        else:
            st.info("No data available")
    
    with tab2:
        df_tech = load_data('SELECT * FROM technical_signals WHERE "Signal" != \'WAIT\' ORDER BY "Score" DESC')
        
        if not df_tech.empty:
            st.dataframe(
                df_tech[[
                    'ticker', 'price', 'rsi', 'Score', 'Signal', 'trigger_type'
                ]].style.format({
                    'price': '${:.2f}',
                    'rsi': '{:.1f}'
                }),
                use_container_width=True,
                height=500,
                hide_index=True
            )
        else:
            st.info("No active signals")
    
    with tab3:
        df_etf = load_data("SELECT sector, weight, hedge_etf, recommendation FROM etf_hedges ORDER BY weight DESC")
        
        if not df_etf.empty:
            st.dataframe(
                df_etf.style.format({"weight": "{:.1%}"}),
                use_container_width=True,
                hide_index=True
            )
        else:
            st.info("No hedge data")
    
    with tab4:
        df_ord = load_data("SELECT * FROM alpha_orders")
        
        if not df_ord.empty:
            st.dataframe(
                df_ord[['Ticker', 'Shares', 'Limit_Price', 'Upside', 'Reason']].style.format({
                    "Limit_Price": "${:.2f}"
                }),
                use_container_width=True,
                hide_index=True
            )
            
            csv = df_ord.to_csv(index=False).encode('utf-8')
            st.download_button(
                "Download CSV",
                csv,
                "orders.csv",
                "text/csv",
                use_container_width=True
            )
        else:
            st.info("No orders generated")

# ============================================================================
# PAGE: RESEARCH
# ============================================================================

def page_research():
    """Research lab - minimal interface"""
    UI.header("Research", "Factor analysis and technical screening")
    
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Factor Scoring", use_container_width=True):
            run_logic_script("model_engine.py")
    with c2:
        if st.button("Technical Sniper", use_container_width=True):
            run_logic_script("technical_models.py")
    
    UI.divider()
    
    t1, t2 = st.tabs(["Rankings", "Signals"])
    
    with t1:
        df_rank = load_data('SELECT * FROM quant_rankings ORDER BY "TOTAL_SCORE" DESC LIMIT 50')
        
        if not df_rank.empty:
            st.dataframe(
                df_rank[[
                    'ticker', 'TOTAL_SCORE', 'current_price', 
                    'pe_ratio', 'profit_margin', 'momentum_pct'
                ]].style.format({
                    'TOTAL_SCORE': '{:.1f}',
                    'current_price': '${:.2f}',
                    'pe_ratio': '{:.1f}',
                    'profit_margin': '{:.2%}',
                    'momentum_pct': '{:.2%}'
                }),
                use_container_width=True,
                height=600,
                hide_index=True
            )
        else:
            st.info("No ranking data")
    
    with t2:
        df_tech = load_data('SELECT * FROM technical_signals ORDER BY "Score" DESC')
        
        if not df_tech.empty:
            # Show full checklist with minimal formatting
            st.dataframe(
                df_tech[[
                    'ticker', 'price', 'Score', 'Signal', 
                    'rsi', 'trend_pass', 'zone_pass', 'trigger_type'
                ]].style.format({
                    'price': '${:.2f}',
                    'rsi': '{:.1f}'
                }),
                use_container_width=True,
                height=600,
                hide_index=True
            )
        else:
            st.info("No technical data")

# ============================================================================
# PAGE: REPORTS
# ============================================================================

def page_reports():
    """AI reports - minimal interface"""
    UI.header("Reports", "AI-powered research analysis")
    
    try:
        from logic.report_engine import generate_ai_report, save_report, get_report_history
    except ImportError:
        st.error("Report engine unavailable")
        return
    
    if "current_report" not in st.session_state:
        st.session_state.current_report = ""
    if "current_ticker" not in st.session_state:
        st.session_state.current_ticker = ""
    
    tab1, tab2 = st.tabs(["Generate", "Archive"])
    
    with tab1:
        c1, c2 = st.columns([1, 3])
        
        with c1:
            db_tickers = load_data("SELECT DISTINCT ticker FROM prices ORDER BY ticker")['ticker'].tolist()
            if not db_tickers:
                db_tickers = ["AAPL", "MSFT", "TSLA", "GOOGL"]
            
            ticker_input = st.selectbox("Ticker", db_tickers, label_visibility="collapsed")
            
            if st.button("Generate", use_container_width=True, type="primary"):
                with st.spinner(f"Analyzing {ticker_input}..."):
                    report = generate_ai_report(ticker_input)
                    st.session_state.current_report = report
                    st.session_state.current_ticker = ticker_input
        
        with c2:
            if st.session_state.current_report:
                st.markdown(st.session_state.current_report)
                
                st.markdown("<br>", unsafe_allow_html=True)
                
                c1, c2, c3 = st.columns(3)
                
                with c1:
                    if st.button("Save", use_container_width=True):
                        version = save_report(
                            st.session_state.current_ticker, 
                            st.session_state.current_report
                        )
                        st.success(f"Saved (v{version})")
                
                with c2:
                    docx_file = create_docx(
                        st.session_state.current_report, 
                        st.session_state.current_ticker
                    )
                    st.download_button(
                        "Word",
                        docx_file,
                        f"{st.session_state.current_ticker}.docx",
                        use_container_width=True
                    )
                
                with c3:
                    pdf_file = create_pdf(
                        st.session_state.current_report, 
                        st.session_state.current_ticker
                    )
                    st.download_button(
                        "PDF",
                        pdf_file,
                        f"{st.session_state.current_ticker}.pdf",
                        use_container_width=True
                    )
    
    with tab2:
        df_hist = get_report_history()
        
        if not df_hist.empty:
            # Minimal archive display
            st.dataframe(
                df_hist[['report_date', 'ticker', 'version', 'signal']],
                use_container_width=True,
                hide_index=True
            )
            
            # Simple viewer
            if len(df_hist) > 0:
                UI.divider()
                report_idx = st.selectbox(
                    "View report",
                    df_hist.index,
                    format_func=lambda x: f"{df_hist.loc[x, 'ticker']} — {df_hist.loc[x, 'report_date']}",
                    label_visibility="collapsed"
                )
                
                if report_idx is not None:
                    st.markdown(df_hist.loc[report_idx, 'content'])
        else:
            st.info("No reports saved")

# ============================================================================
# PAGE: SYSTEM
# ============================================================================

def page_system():
    """System administration"""
    UI.header("System", "Database and diagnostics")
    
    with st.expander("Database", expanded=False):
        c1, c2, c3 = st.columns(3)
        
        with c1:
            if st.button("Reset DB", use_container_width=True):
                run_logic_script("setup_db.py")
        
        with c2:
            if st.button("Ingest Tickers", use_container_width=True):
                run_logic_script("ingest_tickers.py")
        
        with c3:
            if st.button("Update Data", use_container_width=True):
                run_logic_script("data_pipeline.py")
    
    with st.expander("Diagnostics", expanded=False):
        if st.button("Health Check", use_container_width=True):
            run_logic_script("diagnose.py")
        
        # Simple table stats
        if engine:
            tables = ['tickers', 'prices', 'fundamentals', 'alpha_valuation']
            
            stats = []
            for table in tables:
                try:
                    count_df = load_data(f"SELECT COUNT(*) as count FROM {table}")
                    if not count_df.empty:
                        stats.append({
                            'table': table,
                            'records': f"{count_df.iloc[0]['count']:,}"
                        })
                except:
                    pass
            
            if stats:
                st.dataframe(
                    pd.DataFrame(stats),
                    use_container_width=True,
                    hide_index=True
                )

# ============================================================================
# MAIN ROUTER
# ============================================================================

def main():
    """Main application"""
    mode = render_sidebar()
    
    if mode == "Daily Alpha":
        page_daily_alpha()
    elif mode == "Research":
        page_research()
    elif mode == "Reports":
        page_reports()
    elif mode == "System":
        page_system()

if __name__ == "__main__":
    main()
