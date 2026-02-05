import streamlit as st
import pandas as pd
import os
import subprocess
import sys
import time
import plotly.graph_objects as go
from sqlalchemy import text
from logic.db_config import get_engine
from fpdf import FPDF
from docx import Document
from io import BytesIO

# --- CONFIGURATION ---
st.set_page_config(page_title="QuantValue Terminal (Gold)", layout="wide", initial_sidebar_state="expanded")
LOGIC_DIR = "logic"

# --- JONY IVE AESTHETIC (CSS) ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600&family=JetBrains+Mono:wght@400;500&family=Crimson+Pro:wght@400;600&display=swap');
    
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; color: #E0E0E0; }
    h1, h2, h3 { font-family: 'Inter', sans-serif; font-weight: 600; letter-spacing: -0.5px; }
    p, div, li { font-family: 'Crimson Pro', serif; font-size: 1.05rem; color: #B0B0B0; }
    
    .stDataFrame, .stTable, div[data-testid="stMetricValue"] { 
        font-family: 'JetBrains Mono', monospace !important; font-size: 0.95rem; 
    }
    
    div.stButton > button {
        border-radius: 20px; border: 1px solid #404040; background-color: transparent;
        color: #E0E0E0; font-family: 'Inter', sans-serif; font-weight: 600; transition: all 0.2s ease;
    }
    div.stButton > button:hover {
        border-color: #00FF00; color: #00FF00; background-color: rgba(0, 255, 0, 0.05);
    }
    
    div[data-testid="stExpander"] {
        border: 1px solid #303030; border-radius: 6px; background-color: #0e1117;
    }
    
    div[data-testid="stStatusWidget"] {
        border: 1px solid #00FF00; background-color: #051a05;
    }
</style>
""", unsafe_allow_html=True)

# --- DATABASE CONNECTION ---
@st.cache_resource
def get_db_connection():
    try: return get_engine()
    except Exception as e: return None

engine = get_db_connection()

# --- HELPER: EXPORT ENGINE (PDF/DOCX) ---
def create_pdf(text_content, ticker):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # Basic cleanup for latin-1 encoding
    clean_text = text_content.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 10, clean_text)
    
    buffer = BytesIO()
    pdf_output = pdf.output(dest='S').encode('latin-1')
    buffer.write(pdf_output)
    return buffer

def create_docx(text_content, ticker):
    doc = Document()
    doc.add_heading(f'Investment Report: {ticker}', 0)
    for line in text_content.split('\n'):
        doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- HELPER: DOCUMENTATION ENGINE ---
def render_guide(topic):
    if topic == "valuation":
        with st.expander("üìò GUIDE: How to read the Valuation Report"):
            st.markdown("""
            **The Philosophy:** This module calculates the *Intrinsic Value* of a business using a **Discounted Cash Flow (DCF)** model.
            * **Upside %:** The gap between Price and Value. Look for > 15%.
            * **WACC (Risk):** The "Hurdle Rate." Higher = Riskier.
            """)     
    elif topic == "technicals":
        with st.expander("üìò GUIDE: The Sniper's Trinity"):
            st.markdown("""
            **The Philosophy:** We wait for three conditions.
            * **1. The Tide (Trend):** Price > 200 SMA (Bullish).
            * **2. The Zone (Setup):** RSI < 30 (Oversold).
            * **3. The Trigger (Action):** Hammer or Engulfing Candle.
            """)

# --- HELPER: SCRIPT RUNNER ---
def run_logic_script(script_name):
    script_path = os.path.join(LOGIC_DIR, script_name)
    if not os.path.exists(script_path):
        st.error(f"Script not found: {script_path}")
        return
    env = os.environ.copy()
    env["PYTHONPATH"] = os.getcwd()
    st.write(f"üöÄ **Initializing {script_name}...**")
    terminal_placeholder = st.empty()
    logs = []
    process = subprocess.Popen([sys.executable, script_path], stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, env=env, bufsize=1)
    while True:
        line = process.stdout.readline()
        if not line and process.poll() is not None: break
        if line:
            logs.append(line)
            terminal_placeholder.code("".join(logs[-15:]), language="bash")
    if process.returncode == 0: st.success(f"‚úÖ {script_name} Complete.")
    else: st.error(f"‚ùå {script_name} Failed.")

def load_data(query):
    if engine is None: return pd.DataFrame()
    try: return pd.read_sql(query, engine)
    except: return pd.DataFrame()

# --- SIDEBAR ---
st.sidebar.title("QUANT LAB [GOLD]")
st.sidebar.markdown("---")
# REORDERED NAVIGATION
mode = st.sidebar.radio("WORKFLOW", [
    "1. DAILY ACTION (Alpha)", 
    "2. RESEARCH LAB (Deep Dive)", 
    "3. ASSET LAB REPORT (AI)",  # <--- Moved Up
    "‚öôÔ∏è SETTINGS & ADMIN"       # <--- Moved Down & Renamed
])

# ==============================================================================
# MODE 1: DAILY ALPHA
# ==============================================================================
if mode == "1. DAILY ACTION (Alpha)":
    st.title("‚ö° DAILY ALPHA GENERATOR")
    st.markdown("Generate today's buy list based on **Deep Value** and **Price Momentum**.")
    
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1: 
        if st.button("RUN SCANNER"): run_logic_script("scanner.py")
    with c2: 
        if st.button("RUN VALUATION"): run_logic_script("valuation.py")
    with c3: 
        if st.button("RUN CHARTS"): run_logic_script("technical_models.py")
    with c4: 
        if st.button("RUN HEDGE"): run_logic_script("etf_mapper.py")
    with c5: 
        if st.button("GENERATE ORDERS"): run_logic_script("orders.py")

    st.markdown("---")
    
    tab_val, tab_tech, tab_etf, tab_ord = st.tabs(["üíé VALUATION REPORT", "üìà TECHNICAL SIGNALS", "üõ° ETF STRATEGY", "üõí BUY ORDERS"])
    
    with tab_val:
        render_guide("valuation") 
        df = load_data("SELECT * FROM alpha_valuation ORDER BY upside_pct DESC")
        if not df.empty:
            st.metric("Undervalued Opportunities", len(df))
            st.dataframe(df.style.format({"current_price": "${:.2f}", "intrinsic_value": "${:.2f}", "upside_pct": "{:.1f}%", "wacc_pct": "{:.1f}%"}).background_gradient(subset=['upside_pct'], cmap="Greens"), use_container_width=True, height=500)
        else: st.info("No Valuation Data.")

    with tab_tech:
        render_guide("technicals") 
        df_tech = load_data("SELECT * FROM technical_signals WHERE \"Signal\" != 'WAIT' ORDER BY \"Score\" DESC")
        if not df_tech.empty:
             st.dataframe(df_tech.style.applymap(lambda x: 'color: #00FF00' if 'BUY' in str(x) else ('color: #FF0000' if 'SELL' in str(x) else ''), subset=['Signal']), use_container_width=True)
        else: st.info("No Technical Signals.")

    with tab_etf:
        df_etf = load_data("SELECT * FROM etf_hedges ORDER BY weight DESC")
        if not df_etf.empty:
            st.dataframe(df_etf.style.format({"weight": "{:.1%}"}).background_gradient(subset=['weight'], cmap="Reds"), use_container_width=True)

    with tab_ord:
        df_ord = load_data("SELECT * FROM alpha_orders")
        if not df_ord.empty:
            st.dataframe(df_ord.style.format({"Limit_Price": "${:.2f}", "Est_Value": "${:,.2f}"}), use_container_width=True)
            csv = df_ord.to_csv(index=False).encode('utf-8')
            st.download_button("‚¨áÔ∏è Download Order File (CSV)", csv, "basket_orders.csv", "text/csv")

# ==============================================================================
# MODE 2: RESEARCH LAB
# ==============================================================================
elif mode == "2. RESEARCH LAB (Deep Dive)":
    st.title("üî¨ QUANTITATIVE RESEARCH LAB")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("RUN SCORING"): run_logic_script("model_engine.py")
    with c2:
        if st.button("RUN SNIPER"): run_logic_script("technical_models.py")

    st.markdown("---")
    t1, t2 = st.tabs(["üèÜ FACTOR RANKINGS", "üéØ SNIPER SIGNALS"])
    
    with t1:
        df_rank = load_data("SELECT * FROM quant_rankings ORDER BY \"TOTAL_SCORE\" DESC")
        if not df_rank.empty: st.dataframe(df_rank.style.background_gradient(subset=['TOTAL_SCORE'], cmap="RdYlGn"), use_container_width=True, height=600)
    
    with t2:
        render_guide("technicals")
        filter_mode = st.radio("Show:", ["Active Signals Only", "Full Checklist"], horizontal=True)
        df_tech = load_data("SELECT * FROM technical_signals ORDER BY \"Score\" DESC")
        if not df_tech.empty:
            if filter_mode == "Active Signals Only": df_display_source = df_tech[df_tech['Signal'] != 'WAIT']
            else: df_display_source = df_tech
            
            df_display = pd.DataFrame()
            df_display['Ticker'] = df_display_source['ticker']
            df_display['Price'] = df_display_source['price'].apply(lambda x: f"${x:.2f}")
            df_display['Score'] = df_display_source['Score']
            df_display['Signal'] = df_display_source['Signal']
            df_display['üåä Trend'] = df_display_source.apply(lambda x: f"‚úÖ BULL (${x['trend_200_sma']:.2f})" if x['trend_pass'] else f"‚ùå BEAR (${x['trend_200_sma']:.2f})", axis=1)
            df_display['üìâ Zone'] = df_display_source.apply(lambda x: f"‚úÖ LOW ({x['rsi']})" if x['rsi'] < 35 else f"‚ö†Ô∏è ({x['rsi']})", axis=1)
            df_display['üïØÔ∏è Trigger'] = df_display_source['trigger_type'].apply(lambda x: f"‚úÖ {x}" if x != "None" else "‚ùå Wait")
            
            st.dataframe(df_display.style.background_gradient(subset=['Score'], cmap="RdYlGn").map(lambda x: 'color: #00FF00' if 'STRONG' in str(x) else '', subset=['Signal']), use_container_width=True)

# ==============================================================================
# MODE 3: ASSET LAB REPORT (AI)
# ==============================================================================
elif mode == "3. ASSET LAB REPORT (AI)":
    st.title("üß¨ ASSET LAB REPORT")
    st.markdown("Automated Institutional Research powered by **Gemini AI**.")
    st.markdown("---")

    try: 
        from logic.report_engine import generate_ai_report, save_report, get_report_history
    except ImportError: st.error("Report Engine missing.")

    # --- SESSION STATE FOR PERSISTENCE ---
    if "current_report" not in st.session_state: st.session_state.current_report = ""
    if "current_ticker" not in st.session_state: st.session_state.current_ticker = ""

    tab_gen, tab_hist = st.tabs(["‚ú® GENERATOR", "üìú ARCHIVE (Validation)"])

    # --- TAB 1: GENERATOR ---
    with tab_gen:
        c1, c2 = st.columns([1, 3])
        
        with c1:
            st.subheader("New Analysis")
            db_tickers = load_data("SELECT DISTINCT ticker FROM prices ORDER BY ticker")['ticker'].tolist()
            if not db_tickers: db_tickers = ["AAPL", "MSFT", "TSLA"]
            ticker_input = st.selectbox("Ticker Symbol", db_tickers)
            
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button(f"‚ú® GENERATE REPORT", use_container_width=True):
                with st.status(f"üöÄ Analyzing {ticker_input}...", expanded=True) as status:
                    st.write("üì° Fetching Market Data...")
                    time.sleep(0.5)
                    st.write("üß† Synthesizing Narrative...")
                    report = generate_ai_report(ticker_input)
                    st.session_state.current_report = report
                    st.session_state.current_ticker = ticker_input
                    status.update(label="‚úÖ Complete", state="complete", expanded=False)
            
        with c2:
            # Display Report if it exists in session
            if st.session_state.current_report:
                st.markdown(st.session_state.current_report)
                st.markdown("---")
                
                # --- SAVE CONTROLS ---
                sc1, sc2 = st.columns([1, 4])
                with sc1:
                    if st.button("üíæ SAVE TO ARCHIVE"):
                        version = save_report(st.session_state.current_ticker, st.session_state.current_report)
                        st.success(f"Saved to Database! (v{version})")
                
                # --- EXPORT CONTROLS ---
                st.subheader("üìÇ Export")
                ec1, ec2 = st.columns(2)
                docx_file = create_docx(st.session_state.current_report, st.session_state.current_ticker)
                ec1.download_button("üìÑ Download Word", docx_file, f"{st.session_state.current_ticker}_Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
                pdf_file = create_pdf(st.session_state.current_report, st.session_state.current_ticker)
                ec2.download_button("üìï Download PDF", pdf_file, f"{st.session_state.current_ticker}_Report.pdf", "application/pdf")

    # --- TAB 2: ARCHIVE ---
    with tab_hist:
        st.subheader("üìú Research History & Validation")
        
        # Filter Controls
        hist_tickers = get_report_history()['ticker'].unique().tolist()
        if not hist_tickers:
            st.info("No saved reports yet. Generate and Save one!")
        else:
            sel_hist_ticker = st.selectbox("Filter by Ticker", ["All"] + hist_tickers)
            
            # Fetch Data
            df_hist = get_report_history(None if sel_hist_ticker == "All" else sel_hist_ticker)
            
            # Display Table
            st.dataframe(
                df_hist[['report_date', 'ticker', 'version', 'signal', 'created_at']],
                use_container_width=True,
                hide_index=True
            )
            
            # View Selected Report
            st.markdown("---")
            st.write("### üîç Report Viewer")
            report_id = st.selectbox("Select Report to View", df_hist.index, format_func=lambda x: f"{df_hist.loc[x, 'report_date']} - {df_hist.loc[x, 'ticker']} (v{df_hist.loc[x, 'version']})")
            
            if report_id is not None:
                st.markdown(f"**Viewing: {df_hist.loc[report_id, 'ticker']} (v{df_hist.loc[report_id, 'version']})**")
                st.markdown(df_hist.loc[report_id, 'content'])

# ==============================================================================
# MODE 4: SETTINGS & ADMIN
# ==============================================================================
elif mode == "‚öôÔ∏è SETTINGS & ADMIN":
    st.title("‚öôÔ∏è SYSTEM SETTINGS")
    st.warning("‚ö†Ô∏è Restricted Area: Database Operations")
    
    with st.expander("DATABASE MAINTENANCE", expanded=True):
        c1, c2 = st.columns(2)
        with c1: 
            if st.button("RESET DATABASE TABLES"): run_logic_script("setup_db.py")
        with c2: 
            if st.button("INGEST TICKERS"): run_logic_script("ingest_tickers.py")
            
    with st.expander("DATA PIPELINES"):
        if st.button("RUN FULL DATA UPDATE"): run_logic_script("data_pipeline.py")
            
    with st.expander("SYSTEM DIAGNOSTICS"):
        if st.button("RUN HEALTH CHECK"): run_logic_script("diagnose.py")